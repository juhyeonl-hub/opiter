# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 juhyeonl
"""DOCX viewer tab — read-only rendering via python-docx → basic HTML."""
from __future__ import annotations

import html as _html
from pathlib import Path

from PySide6.QtGui import QFont
from PySide6.QtWidgets import QTextEdit, QVBoxLayout

from opiter.ui.cjk_font import cjk_family_chain
from opiter.ui.editors.abstract_editor import AbstractEditor


def _inline(run) -> str:
    text = _html.escape(run.text or "")
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text


def _paragraph_html(para) -> str:
    # ``para.style`` is normally a populated Style object, but DOCX
    # files produced by some editors (or stripped/edited by tooling)
    # leave it as None — defend against that so opening doesn't crash.
    style_obj = getattr(para, "style", None)
    style = (getattr(style_obj, "name", None) or "").lower()
    if style.startswith("heading 1") or style == "title":
        tag = "h1"
    elif style.startswith("heading 2"):
        tag = "h2"
    elif style.startswith("heading 3"):
        tag = "h3"
    elif style.startswith("heading"):
        tag = "h4"
    else:
        tag = "p"
    runs_html = "".join(_inline(r) for r in para.runs)
    if not runs_html:
        runs_html = "&nbsp;"  # preserve blank paragraphs visually
    return f"<{tag}>{runs_html}</{tag}>"


def _table_html(table) -> str:
    rows_html: list[str] = []
    for row in table.rows:
        cells_html: list[str] = []
        for cell in row.cells:
            cell_inner = "".join(_paragraph_html(p) for p in cell.paragraphs)
            cells_html.append(
                f'<td style="border:1px solid #888;padding:4px;">{cell_inner}</td>'
            )
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return (
        '<table style="border-collapse:collapse;margin:6px 0;">'
        + "".join(rows_html)
        + "</table>"
    )


def docx_to_html(path: str | Path) -> str:
    """Render a .docx file as minimal HTML suitable for QTextEdit display.

    Walks body children in document order. For each ``<p>`` / ``<tbl>``
    XML element we pull the next wrapper from ``doc.paragraphs`` /
    ``doc.tables`` — both iterate in document order and only expose
    top-level blocks, which is what we need.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    parts: list[str] = [
        "<html><body style='font-family:sans-serif;line-height:1.4;'>"
    ]
    for block in doc.element.body.iterchildren():
        tag = block.tag.rsplit("}", 1)[-1]
        if tag == "p":
            p = next(para_iter, None)
            if p is not None:
                parts.append(_paragraph_html(p))
        elif tag == "tbl":
            t = next(table_iter, None)
            if t is not None:
                parts.append(_table_html(t))
    parts.append("</body></html>")
    return "".join(parts)


class DOCXEditor(AbstractEditor):
    """Read-only DOCX viewer."""

    def __init__(self) -> None:
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        self._edit = QTextEdit()
        self._edit.setReadOnly(True)
        font = QFont()
        font.setFamilies(cjk_family_chain())
        font.setPointSize(11)
        self._edit.setFont(font)
        layout.addWidget(self._edit)

        self._path: Path | None = None

    # ------------------------------------------------------------ interface
    def open_file(self, path: str | Path) -> None:
        p = Path(path)
        html_str = docx_to_html(p)
        self._edit.setHtml(html_str)
        self._path = p
        self.title_changed.emit(self.display_name())

    def save(self) -> None:
        # Read-only MVP — editing is a later sub-step.
        self.status_message.emit(
            "DOCX editing is not yet supported — use Save As or export to PDF.",
            4000,
        )

    def save_as(self, path: str | Path) -> None:
        self.status_message.emit(
            "DOCX editing is not yet supported.", 4000
        )

    def is_modified(self) -> bool:
        return False

    def file_path(self) -> Path | None:
        return self._path
