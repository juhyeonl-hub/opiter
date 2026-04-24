"""Generate test sample files (DOCX) for manual GUI verification.

Run with: ``uv run python samples/generate_samples.py``

Produces:
- samples/sample.docx — mixed content: headings, body paragraphs,
  bold/italic/underline, lists, a small table.

HWP generation is intentionally not automated — no pure-Python HWP
writer exists. To manually verify the HWP tab you need to supply
your own .hwp file (anything created in Hancom Office or converted
via LibreOffice + h2orestart). Drop it as samples/sample.hwp.
"""
from __future__ import annotations

from pathlib import Path

SAMPLE_DIR = Path(__file__).parent


def make_sample_docx(out: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title + heading hierarchy
    doc.add_heading("Opiter Sample DOCX", level=0)
    doc.add_heading("1. Overview", level=1)
    p = doc.add_paragraph(
        "This document exercises the DOCX viewer tab in "
    )
    run = p.add_run("Opiter")
    run.bold = True
    p.add_run(" — including ")
    p.add_run("italic text").italic = True
    p.add_run(", ")
    p.add_run("underlined text").underline = True
    p.add_run(", and plain text all on one line.")

    doc.add_heading("2. A bulleted list", level=1)
    for item in ("PDF editor with annotations", "DOCX viewer (this tab)",
                 "HWP viewer", "Cross-format export"):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. A numbered list", level=1)
    for i, item in enumerate(
        ("Open a PDF", "Add annotations", "Export as DOCX"), start=1
    ):
        doc.add_paragraph(f"Step {i}: {item}", style="List Number")

    doc.add_heading("4. A small table", level=1)
    table = doc.add_table(rows=4, cols=3)  # 1 header + 3 data rows
    table.style = "Light Grid"
    header = table.rows[0].cells
    header[0].text = "Feature"
    header[1].text = "Status"
    header[2].text = "Notes"
    rows = [
        ("PDF annotation", "Done", "All tools"),
        ("DOCX viewer", "MVP", "Read-only"),
        ("HWP viewer", "MVP", "Text only"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    doc.add_heading("5. Long paragraph for scroll testing", level=1)
    doc.add_paragraph(
        "This paragraph is intentionally verbose so that when you open "
        "the document in Opiter you can test scroll behavior. "
        + ("Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 10)
    )

    # Font sizing nudge so the default render is readable
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    doc.save(str(out))


def main() -> None:
    out = SAMPLE_DIR / "sample.docx"
    make_sample_docx(out)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")
    print("For HWP testing, drop your own .hwp as samples/sample.hwp")


if __name__ == "__main__":
    main()
